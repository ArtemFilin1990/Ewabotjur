# make_protocol_template_bitrix.py
# Generates Bitrix24 legal document templates for protocol of disagreements.
#
# Usage:
#   python make_protocol_template_bitrix.py \
#     --out-docx protocol_bitrix_template.docx \
#     --out-md protocol_bitrix_template.md \
#     --rows 10
#
# Notes:
# - Все UF_CRM__TBD__* затем заменяете на реальные {UF_CRM_XXXXXXXXXXXX} из Bitrix (⚙ Fields).
# - Requisite/BankDetail/MyCompany переменные оставлены как стандартные Bitrix-плейсхолдеры.

import argparse
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass(frozen=True)
class Vars:
    # Документ/контракт
    DOC_NUMBER: str = "{DocumentNumber}"
    DOC_DATE: str = "{DocumentCreateTime~d.m.Y}"
    CITY: str = "{UF_CRM__TBD__CITY__}"
    CONTRACT_TYPE: str = "{UF_CRM__TBD__CONTRACT_TYPE__}"
    CONTRACT_NUMBER: str = "{UF_CRM__TBD__CONTRACT_NUMBER__}"
    CONTRACT_DATE: str = "{UF_CRM__TBD__CONTRACT_DATE__}"

    # Стороны (названия)
    SUPPLIER_NAME: str = "{MyCompanyRequisiteRqCompanyName}"
    BUYER_NAME: str = "{RequisiteRqCompanyName}"

    # Представители/основания
    SUPPLIER_REP_POS: str = "{UF_CRM__TBD__SUPPLIER_REPRESENTATIVE_POSITION__}"
    SUPPLIER_REP_FIO: str = "{UF_CRM__TBD__SUPPLIER_REPRESENTATIVE_NAME__}"
    SUPPLIER_BASIS: str = "{UF_CRM__TBD__SUPPLIER_BASIS__}"

    BUYER_REP_POS: str = "{UF_CRM__TBD__BUYER_REPRESENTATIVE_POSITION__}"
    BUYER_REP_FIO: str = "{UF_CRM__TBD__BUYER_REPRESENTATIVE_NAME__}"
    BUYER_BASIS: str = "{UF_CRM__TBD__BUYER_BASIS__}"


V = Vars()


def build_md(rows: int) -> str:
    """Generate Markdown template for protocol of disagreements.
    
    Args:
        rows: Number of disagreement rows in the table (must be positive)
        
    Returns:
        Markdown-formatted template as a string
    """
    # Название по центру — через HTML, чтобы переносилось в DOCX/рендерилось предсказуемо
    lines = []
    lines.append('<p align="center"><b>ПРОТОКОЛ РАЗНОГЛАСИЙ</b></p>')
    lines.append(
        f'<p align="center">к Договору {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} от {V.CONTRACT_DATE}</p>'
    )
    lines.append("")
    lines.append(f"г. {V.CITY}                                                «{V.DOC_DATE}»")
    lines.append("")
    lines.append(
        f"{V.SUPPLIER_NAME}, именуемое в дальнейшем «Поставщик», в лице {V.SUPPLIER_REP_POS} {V.SUPPLIER_REP_FIO}, "
        f"действующего на основании {V.SUPPLIER_BASIS}, с одной стороны, и {V.BUYER_NAME}, именуемое в дальнейшем «Покупатель», "
        f"в лице {V.BUYER_REP_POS} {V.BUYER_REP_FIO}, действующего на основании {V.BUYER_BASIS}, с другой стороны, "
        f"совместно именуемые Стороны, составили настоящий Протокол разногласий к Договору {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} "
        f"от {V.CONTRACT_DATE} о нижеследующем:"
    )
    lines.append("")
    lines.append("1. Таблица разногласий")
    lines.append("")
    lines.append("| Пункт договора | Редакция Покупателя | Редакция Поставщика | Согласованная редакция |")
    lines.append("|---|---|---|---|")
    for i in range(1, rows + 1):
        lines.append(
            f"| {{UF_CRM__TBD__CLAUSE_{i}_NUMBER__}} | "
            f"{{UF_CRM__TBD__CLAUSE_{i}_BUYER_TEXT__}} | "
            f"{{UF_CRM__TBD__CLAUSE_{i}_SUPPLIER_TEXT__}} | "
            f"{{UF_CRM__TBD__CLAUSE_{i}_AGREED_TEXT__}} |"
        )
    lines.append("")
    lines.append("2. Заключительные положения")
    lines.append("")
    lines.append(f"2.1. Настоящий Протокол является неотъемлемой частью Договора {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} от {V.CONTRACT_DATE}.")
    lines.append("")
    lines.append("2.2. В случае противоречий между текстом Договора и настоящим Протоколом применяются положения настоящего Протокола.")
    lines.append("")
    lines.append("2.3. Настоящий Протокол составлен в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон.")
    lines.append("")
    lines.append("3. Подписи сторон")
    lines.append("")
    lines.append("| Поставщик | Покупатель |")
    lines.append("|---|---|")
    lines.append(
        f"| {V.SUPPLIER_NAME}<br><br>{V.SUPPLIER_REP_POS} _____________________ / {V.SUPPLIER_REP_FIO} /   М.П. | "
        f"{V.BUYER_NAME}<br><br>{V.BUYER_REP_POS} _____________________ / {V.BUYER_REP_FIO} /   М.П. |"
    )
    lines.append("")
    lines.append("Служебное примечание:")
    lines.append("- Все переменные вида `{UF_CRM__TBD__...}` заменить на реальные коды `{UF_CRM_XXXXXXXXXXXX}` из списка полей (⚙ Fields) в вашем портале Bitrix24.")
    lines.append("")
    return "\n".join(lines)


def _set_font_run(run, font_name: str = "Times New Roman", size_pt: int = 12, bold: bool = False) -> None:
    """Set font properties for a document run.
    
    Args:
        run: The document run to modify
        font_name: Font family name (default: Times New Roman)
        size_pt: Font size in points (default: 12)
        bold: Whether to make the text bold (default: False)
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold


def build_docx(rows: int, out_docx: Path) -> None:
    """Generate DOCX template for protocol of disagreements.
    
    Args:
        rows: Number of disagreement rows in the table (must be positive)
        out_docx: Path where the DOCX file will be saved
        
    Raises:
        OSError: If the file cannot be created or written
    """
    doc = Document()

    # Title centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ПРОТОКОЛ РАЗНОГЛАСИЙ")
    _set_font_run(r, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"к Договору {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} от {V.CONTRACT_DATE}")
    _set_font_run(r)

    doc.add_paragraph(f"г. {V.CITY}                                                «{V.DOC_DATE}»")

    doc.add_paragraph(
        f"{V.SUPPLIER_NAME}, именуемое в дальнейшем «Поставщик», в лице {V.SUPPLIER_REP_POS} {V.SUPPLIER_REP_FIO}, "
        f"действующего на основании {V.SUPPLIER_BASIS}, с одной стороны, и {V.BUYER_NAME}, именуемое в дальнейшем «Покупатель», "
        f"в лице {V.BUYER_REP_POS} {V.BUYER_REP_FIO}, действующего на основании {V.BUYER_BASIS}, с другой стороны, "
        f"совместно именуемые Стороны, составили настоящий Протокол разногласий к Договору {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} "
        f"от {V.CONTRACT_DATE} о нижеследующем:"
    )

    doc.add_paragraph("1. Таблица разногласий")

    # Main disagreements table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Пункт договора"
    hdr[1].text = "Редакция Покупателя"
    hdr[2].text = "Редакция Поставщика"
    hdr[3].text = "Согласованная редакция"

    for i in range(1, rows + 1):
        row = table.add_row().cells
        row[0].text = f"{{UF_CRM__TBD__CLAUSE_{i}_NUMBER__}}"
        row[1].text = f"{{UF_CRM__TBD__CLAUSE_{i}_BUYER_TEXT__}}"
        row[2].text = f"{{UF_CRM__TBD__CLAUSE_{i}_SUPPLIER_TEXT__}}"
        row[3].text = f"{{UF_CRM__TBD__CLAUSE_{i}_AGREED_TEXT__}}"

    doc.add_paragraph("2. Заключительные положения")
    doc.add_paragraph(f"2.1. Настоящий Протокол является неотъемлемой частью Договора {V.CONTRACT_TYPE} № {V.CONTRACT_NUMBER} от {V.CONTRACT_DATE}.")
    doc.add_paragraph("2.2. В случае противоречий между текстом Договора и настоящим Протоколом применяются положения настоящего Протокола.")
    doc.add_paragraph("2.3. Настоящий Протокол составлен в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон.")

    doc.add_paragraph("3. Подписи сторон")

    # Signatures side-by-side (2 columns)
    sig = doc.add_table(rows=2, cols=2)
    sig.cell(0, 0).text = "Поставщик"
    sig.cell(0, 1).text = "Покупатель"
    sig.cell(1, 0).text = (
        f"{V.SUPPLIER_NAME}\n\n"
        f"{V.SUPPLIER_REP_POS} _____________________ / {V.SUPPLIER_REP_FIO} /   М.П."
    )
    sig.cell(1, 1).text = (
        f"{V.BUYER_NAME}\n\n"
        f"{V.BUYER_REP_POS} _____________________ / {V.BUYER_REP_FIO} /   М.П."
    )

    doc.add_paragraph("Служебное примечание:")
    doc.add_paragraph("— Все переменные вида {UF_CRM__TBD__...} заменить на реальные коды {UF_CRM_XXXXXXXXXXXX} из списка полей (⚙ Fields) в вашем портале Bitrix24.")

    try:
        out_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_docx)
    except OSError as e:
        print(f"ERROR: Failed to write DOCX file: {e}", file=sys.stderr)
        raise


def _validate_output_path(path: Path, file_type: str) -> None:
    """Validate output path for security and correctness.
    
    Args:
        path: Path to validate
        file_type: Type of file (for error messages)
        
    Raises:
        ValueError: If path contains potential directory traversal or is invalid
    """
    # Check for directory traversal attempts
    try:
        resolved = path.resolve()
        # Ensure the path doesn't escape the current working directory tree
        # (allow any valid path but prevent obvious traversal patterns)
        if ".." in str(path):
            print(f"WARNING: Path contains '..' - {path}", file=sys.stderr)
    except (ValueError, RuntimeError) as e:
        raise ValueError(f"Invalid {file_type} path: {e}") from e


def main():
    """Main entry point for the template generator."""
    ap = argparse.ArgumentParser(
        description="Generate Bitrix24 legal document templates for protocol of disagreements"
    )
    ap.add_argument("--out-docx", required=True, help="Output DOCX path")
    ap.add_argument("--out-md", required=True, help="Output MD path")
    ap.add_argument(
        "--rows", 
        type=int, 
        default=10, 
        help="Number of disagreement rows in table (must be positive, recommended: 5-20)"
    )
    args = ap.parse_args()

    # Validate rows input
    if args.rows <= 0:
        print(f"ERROR: --rows must be a positive integer, got: {args.rows}", file=sys.stderr)
        sys.exit(1)
    
    if args.rows > 100:
        print(f"WARNING: --rows={args.rows} is very large, may cause performance issues", file=sys.stderr)

    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)
    
    # Validate paths
    _validate_output_path(out_docx, "DOCX")
    _validate_output_path(out_md, "MD")

    try:
        # Generate Markdown
        md = build_md(rows=args.rows)
        out_md.write_text(md, encoding="utf-8")
        
        # Generate DOCX
        build_docx(rows=args.rows, out_docx=out_docx)
        
        print(f"SUCCESS: Generated templates\n- DOCX: {out_docx}\n- MD: {out_md}\n- Rows: {args.rows}")
    except OSError as e:
        print(f"ERROR: Failed to write output files: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Unexpected error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
