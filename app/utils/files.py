"""File download and text extraction helpers."""

from __future__ import annotations

import io
import os
from dataclasses import dataclass
from typing import Iterable

from docx import Document
from pypdf import PdfReader


@dataclass(slots=True)
class ExtractedDocument:
    """Extracted document data."""

    text: str
    content_length: int
    filename: str


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _decode_text(raw_bytes: bytes, encodings: Iterable[str]) -> str:
    for encoding in encodings:
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def extract_text_from_bytes(filename: str, content: bytes) -> ExtractedDocument:
    """Extract text from supported document formats."""

    extension = os.path.splitext(filename.lower())[1]
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file extension")

    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        pages_text = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages_text).strip()
    elif extension == ".docx":
        document = Document(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
    else:
        text = _decode_text(content, ["utf-8", "cp1251"]).strip()

    return ExtractedDocument(text=text, content_length=len(content), filename=filename)
