"""
Модуль обработки файлов различных форматов.

Поддержка PDF, DOCX, TXT, MD, изображений с OCR.
"""

from typing import Optional, BinaryIO, Union
from pathlib import Path
from enum import Enum
import io


class FileFormat(Enum):
    """Поддерживаемые форматы файлов."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    IMAGE = "image"


class FileProcessor:
    """Процессор для работы с файлами."""
    
    # Максимальный размер текста (в символах) для защиты от переполнения контекста
    MAX_TEXT_LENGTH = 50000
    
    def __init__(self, enable_ocr: bool = False):
        """
        Инициализация процессора.
        
        Args:
            enable_ocr: Включить OCR для изображений
        """
        self.enable_ocr = enable_ocr
    
    def extract_text(self, file_path: Union[Path, str, io.BytesIO], ext: Optional[str] = None) -> str:
        """
        Извлечение текста из файла.
        
        Args:
            file_path: Путь к файлу или BytesIO буфер
            ext: Расширение файла (с точкой, например ".pdf").
                 Если None, определяется автоматически из file_path.
        
        Returns:
            Извлечённый текст
        
        Raises:
            ValueError: Если формат файла не поддерживается
        """
        # Определяем расширение файла
        if ext is None:
            if isinstance(file_path, (Path, str)):
                ext = Path(file_path).suffix.lower()
            else:
                raise ValueError("Для BytesIO необходимо явно указать расширение ext")
        
        ext = ext.lower()
        
        # Выбираем метод парсинга в зависимости от формата
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        elif ext in [".txt", ".md"]:
            return self.parse_text(file_path)
        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"] and self.enable_ocr:
            return self.ocr_image(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")
    
    def parse_pdf(self, file_path: Union[Path, str, io.BytesIO]) -> str:
        """
        Парсинг PDF файла.
        
        Args:
            file_path: Путь к PDF файлу или BytesIO буфер
        
        Returns:
            Текст из PDF
        """
        try:
            import fitz  # PyMuPDF
            
            text = ""
            # PyMuPDF может работать как с файлами, так и с байтами
            if isinstance(file_path, io.BytesIO):
                doc = fitz.open(stream=file_path.read(), filetype="pdf")
            else:
                doc = fitz.open(str(file_path))
            
            try:
                for page in doc:
                    text += page.get_text()
                    # Проверяем лимит для защиты от слишком больших файлов
                    if len(text) >= self.MAX_TEXT_LENGTH:
                        break
            finally:
                doc.close()
            
            # Обрезаем до максимальной длины
            return text[:self.MAX_TEXT_LENGTH]
            
        except ImportError:
            return "[Ошибка: PyMuPDF не установлен. Установите: pip install PyMuPDF]"
        except Exception as e:
            return f"[Ошибка чтения PDF файла: {e}]"
    
    def parse_docx(self, file_path: Union[Path, str, io.BytesIO]) -> str:
        """
        Парсинг DOCX файла.
        
        Args:
            file_path: Путь к DOCX файлу или BytesIO буфер
        
        Returns:
            Текст из DOCX
        """
        try:
            from docx import Document
            
            # python-docx поддерживает и файлы, и BytesIO
            doc = Document(file_path)
            
            # Извлекаем текст из всех параграфов
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            
            # Обрезаем до максимальной длины
            return text[:self.MAX_TEXT_LENGTH]
            
        except ImportError:
            return "[Ошибка: python-docx не установлен. Установите: pip install python-docx]"
        except Exception as e:
            return f"[Ошибка чтения DOCX файла: {e}]"
    
    def parse_text(self, file_path: Union[Path, str, io.BytesIO]) -> str:
        """
        Чтение обычного текстового файла.
        
        Args:
            file_path: Путь к текстовому файлу или BytesIO буфер
        
        Returns:
            Текст из файла
        """
        try:
            if isinstance(file_path, io.BytesIO):
                text = file_path.read().decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            return text[:self.MAX_TEXT_LENGTH]
            
        except Exception as e:
            return f"[Ошибка чтения текстового файла: {e}]"
    
    def ocr_image(self, file_path: Union[Path, str, io.BytesIO]) -> str:
        """
        OCR распознавание текста на изображении.
        
        Args:
            file_path: Путь к изображению или BytesIO буфер
        
        Returns:
            Распознанный текст
        """
        # TODO: Реализовать OCR с помощью Tesseract или аналога
        # Требует установки pytesseract и Tesseract-OCR
        return "[OCR не реализовано. Требуется установка Tesseract.]"


class DocumentGenerator:
    """Генератор выходных документов."""
    
    def generate_docx(self, template_path: Optional[Path] = None, 
                     context: Optional[dict] = None, 
                     output_path: Optional[Path] = None) -> bytes:
        """
        Генерация DOCX документа.
        
        Args:
            template_path: Путь к шаблону DOCX (для docxtpl)
            context: Данные для подстановки в шаблон
            output_path: Путь для сохранения (опционально)
        
        Returns:
            Байты сгенерированного документа
        
        Raises:
            ImportError: Если docxtpl не установлен
        """
        try:
            from docxtpl import DocxTemplate
            from docx import Document
            import io
            
            if template_path and template_path.exists():
                # Генерация из шаблона
                doc = DocxTemplate(str(template_path))
                doc.render(context or {})
            else:
                # Создание простого документа
                doc = Document()
                if context and 'content' in context:
                    doc.add_paragraph(context['content'])
            
            # Сохранение в BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Опционально сохранить в файл
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(buffer.getvalue())
                buffer.seek(0)
            
            return buffer.getvalue()
            
        except ImportError:
            raise ImportError("docxtpl не установлен. Установите: pip install docxtpl")
        except Exception as e:
            raise Exception(f"Ошибка генерации DOCX: {e}")
    
    def generate_xlsx(self, data: dict, output_path: Optional[Path] = None) -> bytes:
        """
        Генерация XLSX таблицы.
        
        Args:
            data: Данные для таблицы. Может быть словарём с ключами:
                  - 'rows': список списков (данные таблицы)
                  - 'columns': список названий колонок
                  - 'dataframe': pandas DataFrame
            output_path: Путь для сохранения (опционально)
        
        Returns:
            Байты сгенерированной таблицы
        
        Raises:
            ImportError: Если pandas или openpyxl не установлены
        """
        try:
            import pandas as pd
            import io
            
            # Создаём DataFrame из данных
            if 'dataframe' in data:
                df = data['dataframe']
            elif 'rows' in data:
                columns = data.get('columns', None)
                df = pd.DataFrame(data['rows'], columns=columns)
            else:
                # Пытаемся создать DataFrame напрямую из словаря
                df = pd.DataFrame(data)
            
            # Сохранение в BytesIO
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Sheet1')
            
            buffer.seek(0)
            
            # Опционально сохранить в файл
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(buffer.getvalue())
                buffer.seek(0)
            
            return buffer.getvalue()
            
        except ImportError as e:
            if 'pandas' in str(e):
                raise ImportError("pandas не установлен. Установите: pip install pandas")
            elif 'openpyxl' in str(e):
                raise ImportError("openpyxl не установлен. Установите: pip install openpyxl")
            raise
        except Exception as e:
            raise Exception(f"Ошибка генерации XLSX: {e}")


# TODO: Добавить поддержку различных форматов
# TODO: Реализовать OCR с Tesseract
# TODO: Добавить шаблоны для генерации документов
